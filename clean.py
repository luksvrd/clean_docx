from docx import Document
import re

def clean_document(path):
    # Load the existing Word document
    doc = Document(path)
    
    # Create a new Document to save cleaned data
    clean_doc = Document()
    merged_paragraph = clean_doc.add_paragraph()  # Create one single paragraph
    
    # Iterate through each paragraph in the document
    for para in doc.paragraphs:
        text = para.text.strip()  # Strip whitespace from the ends
        
        # Check if the line is not empty and does not contain a timestamp
        if text and not text.startswith('[') and ':' not in text:
            # Append text with a space to merge into one paragraph
            if not re.search(r'[.!?]$', text):
                text += ' '  # Add space for continuous text without punctuation
            else:
                text += ' '  # Add space after punctuation for continuous text
            merged_paragraph.add_run(text)  # Add text to the single paragraph

    # Save the cleaned document
    clean_doc.save('cleaned_document.docx')

# Replace 'path_to_your_document.docx' with the path to your Word document
clean_document('Glover.docx')
